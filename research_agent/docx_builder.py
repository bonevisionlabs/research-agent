"""
Research Agent Toolkit -- DOCX Builder.

Assembles publication-ready Word documents from section text files,
figure images, and inline tables.  Supports an arbitrary number of
papers via a dynamic registration system.

Quick start
-----------
    from pathlib import Path
    from research_agent.docx_builder import (
        PaperConfig, register_paper, compile_paper,
    )

    register_paper(1, PaperConfig(
        title="My First Paper Title",
        authors="Alice Smith, Bob Jones\\nSome University",
        sections_dir=Path("state/paper1/sections"),
        figures_dir=Path("state/paper1/figures"),
        drafts_dir=Path("state/paper1/drafts"),
        figure_captions={"fig1_pipeline": "Figure 1: Overview of ..."},
        tables={
            "table1": {
                "caption": "Table 1: Results summary.",
                "columns": ["Metric", "Value"],
                "rows": [["Accuracy", "0.95"]],
            },
        },
        output_name="Paper_1.docx",
    ))

    compile_paper(paper_id=1)

Convenience wrappers ``compile_paper1()`` and ``compile_paper2()`` are
provided for two-paper projects; they delegate to ``compile_paper()``
and require that the corresponding paper has been registered first.

The document is formatted for academic submission (Times New Roman,
12 pt, double-spaced, first-line indent) and follows standard IMRAD
section ordering.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

# python-docx is imported lazily inside functions that need it,
# so users can import PaperConfig / register_paper without installing docx.
def _import_docx():
    """Lazy import of python-docx (raises ImportError with helpful message)."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise ImportError(
            "python-docx is required for document compilation. "
            "Install it with: pip install research-agent[full] or pip install python-docx"
        )
    return Document, Inches, Pt, Cm, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT

from . import config

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Section ordering -- standard IMRAD format (Scientific Reports style)
# ---------------------------------------------------------------------------

PAPER_SECTIONS: List[str] = [
    "abstract",
    "introduction",
    "methods",
    "results",
    "discussion",
    "conclusion",
    "data_availability",
    "ethics",
    "author_contributions",
    "competing_interests",
    "acknowledgements",
    "references",
]

SECTION_TITLES: Dict[str, str] = {
    "abstract": "Abstract",
    "introduction": "Introduction",
    "methods": "Methods",
    "results": "Results",
    "discussion": "Discussion",
    "conclusion": "Conclusion",
    "data_availability": "Data Availability",
    "ethics": "Ethics Declaration",
    "author_contributions": "Author Contributions",
    "competing_interests": "Competing Interests",
    "acknowledgements": "Acknowledgements",
    "references": "References",
}

# ---------------------------------------------------------------------------
# Regex patterns for inline markers in section text
# ---------------------------------------------------------------------------

_SUBSECTION_RE = re.compile(r"^\d+\.\d+\s+\S")
_FIGURE_RE = re.compile(r"^\[FIGURE:\s*(\S+)\]\s*$")
_FIGURE_WITH_CAPTION_RE = re.compile(r"^\[FIGURE:\s*(\S+)\s*--\s*(.+)\]\s*$")
_TABLE_RE = re.compile(r"^\[TABLE\s+\d+:")

# ---------------------------------------------------------------------------
# PaperConfig dataclass
# ---------------------------------------------------------------------------

TableSpec = Dict[str, Any]
"""A single table specification dict with keys 'caption', 'columns', 'rows'."""


@dataclass
class PaperConfig:
    """Configuration for a single paper managed by the DOCX builder.

    Parameters
    ----------
    title:
        Full paper title (appears centered at top of document).
    authors:
        Author line.  Use ``\\n`` to separate name line from affiliation.
    sections_dir:
        Directory containing ``{section_name}.txt`` files.
    figures_dir:
        Directory containing figure images (``fig*.png``).
    drafts_dir:
        Directory where compiled DOCX files are saved.
    figure_captions:
        Mapping of ``{figure_stem: caption_text}`` for inline figure
        placement.  The stem must match the filename without extension
        (e.g., ``"fig1_pipeline"`` for ``fig1_pipeline.png``).
    tables:
        Mapping of ``{table_key: table_spec}`` for inline table rendering.
        Each *table_spec* is a dict with keys ``"caption"`` (str),
        ``"columns"`` (list[str]), and ``"rows"`` (list[list[str]] or
        list[dict]).  The *table_key* should follow the pattern
        ``"tableN"`` (e.g., ``"table1"``, ``"table2"``).
    output_name:
        Default filename for the compiled DOCX (e.g., ``"Paper_1.docx"``).
    """

    title: str = "Untitled Paper"
    authors: str = ""
    sections_dir: Optional[Path] = None
    figures_dir: Optional[Path] = None
    drafts_dir: Optional[Path] = None
    figure_captions: Dict[str, str] = field(default_factory=dict)
    tables: Dict[str, TableSpec] = field(default_factory=dict)
    output_name: str = "paper.docx"


# ---------------------------------------------------------------------------
# Paper registry
# ---------------------------------------------------------------------------

_registry: Dict[int, PaperConfig] = {}


def register_paper(paper_id: int, paper_config: PaperConfig) -> None:
    """Register (or replace) a paper configuration.

    Parameters
    ----------
    paper_id:
        Integer identifier for the paper (e.g., 1, 2, 3).
    paper_config:
        A :class:`PaperConfig` instance with all paper-specific settings.

    Example
    -------
    ::

        register_paper(1, PaperConfig(
            title="My Paper",
            authors="Alice, Bob",
            sections_dir=Path("sections/"),
            figures_dir=Path("figures/"),
            drafts_dir=Path("drafts/"),
        ))
    """
    _registry[paper_id] = paper_config
    logger.debug("Registered paper %d: %s", paper_id, paper_config.title)


def get_paper_config(paper_id: int) -> PaperConfig:
    """Retrieve a registered paper configuration.

    Raises
    ------
    KeyError
        If *paper_id* has not been registered via :func:`register_paper`.
    """
    try:
        return _registry[paper_id]
    except KeyError:
        registered = sorted(_registry) or "none"
        raise KeyError(
            f"Paper {paper_id} is not registered. "
            f"Call register_paper({paper_id}, PaperConfig(...)) first. "
            f"Currently registered: {registered}"
        ) from None


def registered_papers() -> Dict[int, PaperConfig]:
    """Return a copy of the current paper registry."""
    return dict(_registry)


def clear_registry() -> None:
    """Remove all registered papers (useful in tests)."""
    _registry.clear()


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _resolve_paper_dirs(
    paper_id: int,
    pc: PaperConfig,
) -> Dict[str, Any]:
    """Build a resolved directory dict for a paper.

    If the PaperConfig has explicit paths, use them.  Otherwise, try to
    derive paths from the module-level ``config.ProjectConfig`` (if the
    paper was registered there too).
    """
    # Start with explicit PaperConfig paths
    sections_dir = pc.sections_dir
    figures_dir = pc.figures_dir
    drafts_dir = pc.drafts_dir

    # Fall back to ProjectConfig paths if available
    if sections_dir is None or figures_dir is None or drafts_dir is None:
        try:
            cfg = config.get_config()
            pp = cfg.paper(paper_id)
            if sections_dir is None:
                sections_dir = pp.sections
            if figures_dir is None:
                figures_dir = pp.figures
            if drafts_dir is None:
                drafts_dir = pp.drafts
        except (KeyError, AttributeError):
            pass

    return {
        "title": pc.title,
        "authors": pc.authors,
        "sections_dir": sections_dir,
        "figures_dir": figures_dir,
        "drafts_dir": drafts_dir,
        "figure_captions": pc.figure_captions,
        "tables": pc.tables,
        "output_name": pc.output_name,
    }


def _load_sections(sections_dir: Path) -> Dict[str, str]:
    """Load section text files from a directory.

    Looks for files named ``{section_name}.txt`` for each entry in
    :data:`PAPER_SECTIONS`.  Missing files are silently skipped.
    """
    sections: Dict[str, str] = {}
    for section_name in PAPER_SECTIONS:
        path = sections_dir / f"{section_name}.txt"
        if path.exists():
            sections[section_name] = path.read_text(encoding="utf-8")
    return sections


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------

def compile_paper(
    paper_id: int = 1,
    sections: Optional[Dict[str, str]] = None,
    title: Optional[str] = None,
    authors: Optional[str] = None,
    output_path: Optional[Path] = None,
    include_figures: bool = True,
) -> Path:
    """Compile a registered paper into a publication-ready DOCX document.

    The paper must have been registered via :func:`register_paper` before
    calling this function.

    Parameters
    ----------
    paper_id:
        Integer identifier for the paper to compile.
    sections:
        Dict of ``{section_name: text}``.  If ``None``, sections are
        loaded from the paper's ``sections_dir``.
    title:
        Paper title override.  If ``None``, uses the registered title.
    authors:
        Author line override.  If ``None``, uses the registered authors.
    output_path:
        Where to save the DOCX.  If ``None``, saves to the paper's
        ``drafts_dir`` with the registered ``output_name``.
    include_figures:
        Whether to insert inline figures (default ``True``).

    Returns
    -------
    Path
        Resolved path to the saved DOCX file.

    Raises
    ------
    KeyError
        If *paper_id* has not been registered.
    ValueError
        If no sections directory is configured and *sections* is not
        provided.
    """
    pc = get_paper_config(paper_id)
    resolved = _resolve_paper_dirs(paper_id, pc)

    if title is None:
        title = resolved["title"]
    if authors is None:
        authors = resolved["authors"]

    # Ensure output directories exist
    try:
        cfg = config.get_config()
        cfg.ensure_dirs()
    except Exception:
        pass

    Document, Inches, Pt, Cm, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT = _import_docx()

    doc = Document()

    # -- Styles ---------------------------------------------------------------
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0)

    # -- Title ----------------------------------------------------------------
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = "Times New Roman"

    # -- Authors --------------------------------------------------------------
    authors_para = doc.add_paragraph()
    authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors_run = authors_para.add_run(authors)
    authors_run.font.size = Pt(11)
    authors_run.font.name = "Times New Roman"

    doc.add_paragraph()  # spacer

    # -- Load sections --------------------------------------------------------
    if sections is None:
        sections_dir = resolved["sections_dir"]
        if sections_dir is None:
            raise ValueError(
                f"Paper {paper_id} has no sections_dir configured and no "
                f"sections dict was provided.  Either set sections_dir in "
                f"PaperConfig or pass sections= to compile_paper()."
            )
        sections = _load_sections(Path(sections_dir))

    # -- Add sections ---------------------------------------------------------
    figures_dir = resolved["figures_dir"]
    figure_captions = resolved["figure_captions"]
    tables_dict = resolved["tables"]

    for section_name in PAPER_SECTIONS:
        heading_text = SECTION_TITLES.get(section_name, section_name.title())
        heading = doc.add_heading(heading_text, level=1)
        heading.runs[0].font.name = "Times New Roman"
        heading.runs[0].font.size = Pt(13)

        text = sections.get(section_name, "")
        if text:
            for paragraph_text in text.split("\n\n"):
                paragraph_text = paragraph_text.strip()
                if not paragraph_text:
                    continue

                # -- Inline figure marker: [FIGURE: name] or [FIGURE: name -- caption]
                fig_match = (
                    _FIGURE_RE.match(paragraph_text)
                    or _FIGURE_WITH_CAPTION_RE.match(paragraph_text)
                )
                if fig_match and include_figures and figures_dir is not None:
                    fig_name = fig_match.group(1)
                    fig_path = Path(figures_dir) / f"{fig_name}.png"
                    caption = figure_captions.get(
                        fig_name, f"Figure: {fig_name}"
                    )
                    if fig_path.exists():
                        try:
                            insert_figure(doc, fig_path, caption)
                        except Exception as e:
                            logger.warning(
                                "Failed to insert inline figure %s: %s",
                                fig_name, e,
                            )

                # -- Inline table marker: [TABLE N: ...]
                elif _TABLE_RE.match(paragraph_text):
                    inserted = False
                    for table_key, table_data in tables_dict.items():
                        tnum = table_key.replace("table", "")
                        if f"[TABLE {tnum}:" in paragraph_text:
                            insert_table(
                                doc,
                                table_data["columns"],
                                table_data["rows"],
                                caption=table_data.get("caption", ""),
                            )
                            inserted = True
                            break
                    if not inserted:
                        para = doc.add_paragraph(paragraph_text)
                        para.paragraph_format.first_line_indent = Cm(1.27)

                # -- Markdown-style headings
                elif paragraph_text.startswith("##"):
                    sub_heading = paragraph_text.lstrip("#").strip()
                    h = doc.add_heading(sub_heading, level=2)
                    h.runs[0].font.name = "Times New Roman"
                    h.runs[0].font.size = Pt(12)
                elif paragraph_text.startswith("#"):
                    sub_heading = paragraph_text.lstrip("#").strip()
                    h = doc.add_heading(sub_heading, level=2)
                    h.runs[0].font.name = "Times New Roman"

                # -- Numbered subsection (e.g., "2.1 Preprocessing")
                elif _SUBSECTION_RE.match(paragraph_text):
                    h = doc.add_heading(paragraph_text, level=2)
                    h.runs[0].font.name = "Times New Roman"
                    h.runs[0].font.size = Pt(12)

                # -- Normal paragraph
                else:
                    para = doc.add_paragraph(paragraph_text)
                    para.paragraph_format.first_line_indent = Cm(1.27)
        else:
            doc.add_paragraph(
                f"[{section_name.upper()} section to be completed]",
                style="Normal",
            )

    # -- Save -----------------------------------------------------------------
    if output_path is None:
        drafts_dir = resolved["drafts_dir"]
        if drafts_dir is None:
            raise ValueError(
                f"Paper {paper_id} has no drafts_dir configured and no "
                f"output_path was provided."
            )
        output_path = Path(drafts_dir) / resolved["output_name"]

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    logger.info("Compiled paper %d: %s", paper_id, output_path)
    return output_path


# ---------------------------------------------------------------------------
# Convenience wrappers for two-paper projects
# ---------------------------------------------------------------------------

def compile_paper1(**kwargs) -> Path:
    """Compile paper 1.

    Requires that paper 1 has been registered via :func:`register_paper`.
    All keyword arguments are forwarded to :func:`compile_paper`.
    """
    return compile_paper(paper_id=1, **kwargs)


def compile_paper2(**kwargs) -> Path:
    """Compile paper 2.

    Requires that paper 2 has been registered via :func:`register_paper`.
    All keyword arguments are forwarded to :func:`compile_paper`.
    """
    return compile_paper(paper_id=2, **kwargs)


# ---------------------------------------------------------------------------
# Figure insertion
# ---------------------------------------------------------------------------

def insert_figure(doc: "Document", figure_path: Path, caption: str) -> None:
    """Insert a figure with caption into the document.

    Parameters
    ----------
    doc:
        The python-docx ``Document`` to modify.
    figure_path:
        Path to the image file (PNG, JPEG, etc.).
    caption:
        Caption text rendered below the image in 10 pt italic.
    """
    _, Inches, Pt, _, WD_ALIGN_PARAGRAPH, _ = _import_docx()

    doc.add_paragraph()
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_para.add_run()
    run.add_picture(str(figure_path), width=Inches(6.0))

    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(caption)
    cap_run.font.size = Pt(10)
    cap_run.font.italic = True
    cap_run.font.name = "Times New Roman"


# ---------------------------------------------------------------------------
# Table insertion
# ---------------------------------------------------------------------------

def insert_table(
    doc: "Document",
    columns: List[str],
    rows: List[Any],
    caption: str = "",
) -> None:
    """Insert a formatted table into the document.

    Parameters
    ----------
    doc:
        The python-docx ``Document`` to modify.
    columns:
        List of column header strings.
    rows:
        List of rows.  Each row can be a ``list`` (positional) or a
        ``dict`` keyed by column name.
    caption:
        Optional caption rendered above the table in 10 pt bold.
    """
    if not columns or not rows:
        return

    _, _, Pt, _, _, WD_TABLE_ALIGNMENT = _import_docx()

    if caption:
        cap_para = doc.add_paragraph()
        cap_run = cap_para.add_run(caption)
        cap_run.bold = True
        cap_run.font.size = Pt(10)
        cap_run.font.name = "Times New Roman"

    table = doc.add_table(rows=1 + len(rows), cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, col_name in enumerate(columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col_name)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "Times New Roman"

    # Data rows
    for i, row_data in enumerate(rows):
        for j, col_name in enumerate(columns):
            cell = table.rows[i + 1].cells[j]
            if isinstance(row_data, dict):
                value = row_data.get(col_name, "")
            elif isinstance(row_data, (list, tuple)) and j < len(row_data):
                value = row_data[j]
            else:
                value = ""
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Times New Roman"


# ---------------------------------------------------------------------------
# Append remaining figures (utility)
# ---------------------------------------------------------------------------

def append_figures(
    doc: "Document",
    figures_dir: Path,
    figure_captions: Optional[Dict[str, str]] = None,
) -> None:
    """Append all ``fig*.png`` images from a directory to the document.

    Useful for placing figures at the end of the manuscript (e.g., for
    journals that require figures after references).

    Parameters
    ----------
    doc:
        The python-docx ``Document`` to modify.
    figures_dir:
        Directory to scan for ``fig*.png`` files.
    figure_captions:
        Optional mapping of ``{stem: caption}``.  Falls back to
        ``"Figure: {stem}"`` for unrecognised files.
    """
    if not Path(figures_dir).is_dir():
        return

    figure_files = sorted(Path(figures_dir).glob("fig*.png"))
    if not figure_files:
        return

    if figure_captions is None:
        figure_captions = {}

    doc.add_page_break()
    heading = doc.add_heading("Figures", level=1)
    heading.runs[0].font.name = "Times New Roman"

    for fig_path in figure_files:
        caption = figure_captions.get(fig_path.stem, f"Figure: {fig_path.stem}")
        try:
            insert_figure(doc, fig_path, caption)
        except Exception as e:
            logger.warning("Failed to insert figure %s: %s", fig_path, e)


# ---------------------------------------------------------------------------
# Section persistence
# ---------------------------------------------------------------------------

def save_section(section_name: str, text: str, paper_id: int = 1) -> Path:
    """Save a section's text to the registered paper's sections directory.

    Parameters
    ----------
    section_name:
        Section key (e.g., ``"abstract"``, ``"methods"``).
    text:
        Section content to write.
    paper_id:
        Which registered paper to save under.

    Returns
    -------
    Path
        Path to the saved text file.

    Raises
    ------
    KeyError
        If *paper_id* has not been registered.
    ValueError
        If no sections directory is configured for the paper.
    """
    pc = get_paper_config(paper_id)
    resolved = _resolve_paper_dirs(paper_id, pc)
    sections_dir = resolved["sections_dir"]

    if sections_dir is None:
        raise ValueError(
            f"Paper {paper_id} has no sections_dir configured.  "
            f"Set sections_dir in PaperConfig or register the paper "
            f"in ProjectConfig."
        )

    sections_dir = Path(sections_dir)
    sections_dir.mkdir(parents=True, exist_ok=True)
    path = sections_dir / f"{section_name}.txt"
    path.write_text(text, encoding="utf-8")
    return path
